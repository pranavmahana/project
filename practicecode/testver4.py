import tkinter as tk
from tkinter import filedialog, messagebox, Entry, StringVar, OptionMenu
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.widgets import SpanSelector
from matplotlib.ticker import EngFormatter
from scipy.signal import welch
from PIL import Image
from docx import Document
from docx.shared import Inches
from io import BytesIO

class PSDPlotterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PSD Plotter")
        self.file_path = None
        self.df = None
        self.selected_range = None
        self.sensitivity_var = None
        self.sensitivity_entry = None
        self.fig = None
        self.ax = None
        self.canvas = None
        self.span_selector = None
        self.g_levels_data = None

        self.create_widgets()

    def create_widgets(self):
        # File Selection Button
        self.select_file_button = tk.Button(self.root, text="Select CSV File", command=self.load_file)
        self.select_file_button.pack(pady=10)

        # Sensitivity Checkbox
        self.sensitivity_var = tk.IntVar()
        self.sensitivity_checkbox = tk.Checkbutton(self.root, text="Enter Sensor Sensitivity", variable=self.sensitivity_var, command=self.toggle_sensitivity_entry, bg="#f0f8ff")
        self.sensitivity_checkbox.pack()

        # Entry for Sensitivity
        self.sensitivity_label = tk.Label(self.root, text="Enter Sensor Sensitivity (Hz/V):", bg="#f0f8ff")
        self.sensitivity_label.pack()
        self.sensitivity_entry = Entry(self.root, state="disabled")
        self.sensitivity_entry.pack()

        # Plot G-Levels Button
        self.plot_g_levels_button = tk.Button(self.root, text="Plot G-Levels", command=self.plot_g_levels, bg="#87ceeb")
        self.plot_g_levels_button.pack(pady=10)

        # Plot PSD Button
        self.plot_button = tk.Button(self.root, text="Plot PSD", command=self.plot_psd, bg="#87ceeb")
        self.plot_button.pack(pady=10)

        # Matplotlib Figure
        self.fig, self.ax = plt.subplots(figsize=(8, 6))
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.root)
        self.canvas.get_tk_widget().pack()

        # Span Selector
        self.span_selector = SpanSelector(self.ax, self.onselect, 'horizontal', useblit=True)




        # Export to Word Button
        self.export_button = tk.Button(self.root, text="Export to Word", command=self.export_to_docx, bg="#87ceeb")
        self.export_button.pack(pady=10)

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])

        if file_path:
            self.file_path = file_path
            self.df = pd.read_csv(file_path)
            tk.messagebox.showinfo("File Loaded", "CSV file loaded successfully.")

    def toggle_sensitivity_entry(self):
        if self.sensitivity_var.get() == 1:
            self.sensitivity_entry.config(state="normal")
        else:
            self.sensitivity_entry.config(state="disabled")
            self.sensitivity_entry.delete(0, tk.END)  # Clear the entry

    def plot_g_levels(self):
        if self.df is None:
            tk.messagebox.showerror("Error", "Please select a CSV file first.")
            return

        # Assuming second column contains voltage readings
        voltage_col = self.df.iloc[:, 1]

        # Get Sensitivity from the user input
        if self.sensitivity_var.get() == 1:
            sensitivity_str = self.sensitivity_entry.get()
            try:
                sensitivity = float(sensitivity_str)
            except ValueError:
                tk.messagebox.showerror("Error", "Please enter a valid numeric value for sensitivity.")
                return
        else:
            tk.messagebox.showerror("Error", "Please enter the sensitivity value.")
            return

        # Convert voltage to g-levels
        g_levels = voltage_col / sensitivity

        # Plot G-Levels
        self.ax.clear()
        self.ax.plot(self.df.iloc[:, 0], g_levels, label='G-Levels')
        self.ax.set_xlabel('Time')
        self.ax.set_ylabel('G-Levels')
        self.ax.set_title('G-Levels Over Time')
        self.ax.legend()

        self.canvas.draw()

        # Save g-levels for later use
        self.g_levels_data = pd.DataFrame({'Time': self.df.iloc[:, 0], 'G-Levels': g_levels})

        # Span Selector
        self.span_selector = SpanSelector(self.ax, self.onselect, 'horizontal', useblit=True)

    def plot_psd(self, start_index=None, end_index=None):
        if self.df is None:
            tk.messagebox.showerror("Error", "Please select a CSV file first.")
            return

        # Assuming second column contains voltage readings
        voltage_col = self.df.iloc[:, 1]

        # Sensitivity factor in Hz/V (replace with the actual sensitivity value)
        sensitivity = float(self.sensitivity_entry.get())

        # Convert voltage to g-levels
        g_levels = voltage_col / sensitivity

        if start_index is not None and end_index is not None:
            # Select data within the specified span
            selected_data = self.g_levels_data[(self.g_levels_data['Time'] >= start_index) &
                                               (self.g_levels_data['Time'] <= end_index)]

            if selected_data.empty:
                tk.messagebox.showwarning("No Data", "No data available within the selected span.")
                return

            # Plot PSD for the selected span
            fs = 1.0  # Adjust this value based on your data
            n = len(selected_data['G-Levels'])

            # Get additional parameters for welch function
            try:
                nperseg = int(self.nperseg_entry.get())
                noverlap = int(self.noverlap_entry.get())
                nfft = int(self.nfft_entry.get())
            except ValueError:
                tk.messagebox.showerror("Error", "Please enter valid numeric values for welch parameters.")
                return

            window_type = self.window_type_var.get()

            f, Pxx = self.calculate_psd(selected_data['G-Levels'], fs, n, nperseg=nperseg, noverlap=noverlap, nfft=nfft, window_type=window_type)
        else:
            # Plot PSD for the entire data
            fs = 1.0  # Adjust this value based on your data
            n = len(g_levels)

            # Get additional parameters for welch function
            try:
                nperseg = int(self.nperseg_entry.get())
                noverlap = int(self.noverlap_entry.get())
                nfft = int(self.nfft_entry.get())
            except ValueError:
                tk.messagebox.showerror("Error", "Please enter valid numeric values for welch parameters.")
                return

            window_type = self.window_type_var.get()

            f, Pxx = self.calculate_psd(g_levels, fs, n, nperseg=nperseg, noverlap=noverlap, nfft=nfft, window_type=window_type)

        # Plot PSD
        self.ax.clear()
        self.ax.semilogy(f, Pxx)
        self.ax.set_xlabel('Frequency [Hz]')
        self.ax.set_ylabel('Power/Frequency [dB/Hz]')
        self.ax.set_title('Power Spectral Density')

        # Set the formatter for the y-axis to display engineering notation
        self.ax.yaxis.set_major_formatter(EngFormatter(unit='dB'))

        self.canvas.draw()

    def onselect(self, xmin, xmax):
        if xmin != xmax:
            self.plot_psd(xmin, xmax)

    def export_to_docx(self):
        if self.fig:
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
            if file_path:
                document = Document()
                document.add_heading('PSD Plot and G-Levels Plot', 0)

                # Export G-Levels plot
                g_levels_img_data = BytesIO()
                self.fig.savefig(g_levels_img_data, format='png', dpi=300)
                g_levels_img = Image.open(g_levels_img_data)
                g_levels_img_data.close()
                g_levels_img_path = file_path.replace('.docx', '_g_levels_plot.png')
                g_levels_img.save(g_levels_img_path)
                document.add_heading('G-Levels Plot', level=1)
                document.add_picture(g_levels_img_path, width=Inches(5))

                # Export PSD plot
                psd_img_data = BytesIO()
                self.fig.savefig(psd_img_data, format='png', dpi=300)
                psd_img = Image.open(psd_img_data)
                psd_img_data.close()
                psd_img_path = file_path.replace('.docx', '_psd_plot.png')
                psd_img.save(psd_img_path)
                document.add_heading('PSD Plot (Selected Span)', level=1)
                document.add_picture(psd_img_path, width=Inches(5))

                document.save(file_path)

                tk.messagebox.showinfo("Export Successful", "Plots exported successfully.")
        else:
            tk.messagebox.showerror("Error", "Please plot the PSD before exporting.")

if __name__ == "__main__":
    root = tk.Tk()
    app = PSDPlotterApp(root)
    root.mainloop()
